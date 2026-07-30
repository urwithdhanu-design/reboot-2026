"""Generate Canton layer assessment Word document. Run: python docs/generate_canton_assessment_docx.py"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "CANTON-LAYER-ASSESSMENT.docx"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    title = doc.add_heading("GCUL / Reboot 2026 — Canton Layer Assessment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Reference document · Generated {date.today().isoformat()}\n").bold = True
    meta.add_run(
        "Scope: Daml sandbox, blockchain-orchestrator, policy/claims integration, local dev, and gaps vs enterprise Canton (e.g. LBG commercial banking)."
    )

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "The platform has a working demo path: Daml sandbox → blockchain-orchestrator → policy mint → claims verify, "
        "with a simulated fallback when Canton is unavailable. This is appropriate for Reboot 2026 demos, but several "
        "behaviours can present as “Canton verified” while ledger proof is weak or missing. Production readiness requires "
        "honest mint modes, correct Daml exercise paths, security hardening, and observability beyond orchestrator :8088."
    )

    doc.add_heading("1. Architecture today", level=1)
    doc.add_paragraph(
        "policy-service → blockchain-orchestrator (:8088)\n"
        "  • Canton JSON API (:7575 local / :8080 Cloud Run) → InsurancePolicy / ClaimSettlement\n"
        "  • SimulatedLedgerAdapter when Canton is down\n"
        "  • H2 insurance chain (audit mirror, not Daml)\n"
        "  • gcul-sidecar (:8091) for GBP wallet movements (not on Canton)\n\n"
        "claims-service → verify + settle via orchestrator\n"
        "admin UI → Tokenization / Smart Contracts (live status); capital-market pages are design reference"
    )

    ports = [
        ["Canton JSON Ledger API", "7575 (local)", "http://127.0.0.1:7575"],
        ["blockchain-orchestrator", "8088", "Mint, verify, settle, Canton status"],
        ["policy-service", "8082", "Issues policies, calls mint API"],
        ["claims-service", "8085", "Canton verify on claim intake"],
        ["gcul-sidecar", "8091", "GCUL Universal Ledger (GBP, not Canton)"],
    ]
    doc.add_heading("Key ports (local)", level=2)
    add_table(doc, ["Component", "Port", "Role"], ports)

    doc.add_heading("2. Current problems", level=1)

    doc.add_heading("2.1 Integrity — “minted on Canton” can be misleading", level=2)
    add_bullets(
        doc,
        [
            "Silent simulated fallback: if Canton is down, LedgerAdapterRegistry mints via SimulatedLedgerAdapter. Policy still shows MINTED with a synthetic token_id.",
            "Claims verify bypass: if Canton verify fails but mint_status=MINTED and token_id exists, claims still pass (BlockchainValidationClient.assertVerifiedOnCanton).",
            "Verify uses primary Canton adapter only; simulated mints fail verify on orchestrator, then claims bypass via the rule above.",
            "Impact: customer evaluation trail can show Canton verification passed without a Daml contract.",
        ],
    )

    doc.add_heading("2.2 Daml contract vs Java implementation mismatch", level=2)
    doc.add_paragraph(
        "Daml defines InsurerMintAuthority.MintPolicy as the intended mint gate (Gcul/InsurancePolicy.daml). "
        "Java CantonJsonApiClient.mintPolicy() calls POST /v1/create on InsurancePolicy directly and sends mintedAt as an ISO string; "
        "Daml expects Time. findAuthorityContractId() exists but is not used in the mint path."
    )
    add_bullets(
        doc,
        [
            "Works on permissive sandbox; may fail on stricter Canton Network deployments.",
            "Hardcoded Daml package ID in application.properties drifts on every Daml rebuild.",
        ],
    )

    doc.add_heading("2.3 Configuration and race conditions", level=2)
    add_bullets(
        doc,
        [
            "Local JSON API :7575 vs Cloud Run Canton :8080 — env must match target.",
            "Double mint risk: policy-service may publish PolicyMintRequested and call mint HTTP immediately in cloud.",
            "local-dev.cmd can start APIs before Canton is healthy → first mints silently simulate.",
        ],
    )

    doc.add_heading("2.4 Security (not production-ready)", level=2)
    add_bullets(
        doc,
        [
            "Internal /api/blockchain/internal/* routes lack service-to-service authentication.",
            "Canton sandbox uses unsigned JWTs and --allow-insecure-tokens.",
            "Cloud deploy may expose gcul-canton with allow-unauthenticated.",
            "On-ledger fields include customerId, walletAddress, metadataUri — privacy model not documented.",
        ],
    )

    doc.add_heading("2.5 Settlement and observability", level=2)
    add_bullets(
        doc,
        [
            "GBP payouts use sidecar/wallet; ClaimSettlement on Canton is best-effort (deferred on failure).",
            "Claim settlement may send amountGbp as string vs Daml Decimal.",
            "Observability monitors orchestrator :8088, not Canton :7575.",
            "No metric for simulated fallback rate or package ID mismatch.",
            "Canton Docker sandbox is ephemeral — restart wipes ledger while policy DB still says MINTED.",
        ],
    )

    doc.add_heading("2.6 vs enterprise Canton (LBG direction)", level=2)
    lbg_rows = [
        ["Tokenised deposit + DvP + bank reconciliation", "Mint + verify; weak proof when simulated"],
        ["Own validator node, bank controls", "Single sandbox, insurer admin JWT"],
        ["Privacy by sub-transaction", "Policy payload on template; customer observer only"],
        ["Oracle-governed triggers", "Parametric off-chain; cat triggers not on Daml"],
    ]
    add_table(doc, ["LBG-style production", "This stack today"], lbg_rows)

    doc.add_heading("3. How to improve (prioritized)", level=1)

    doc.add_heading("Phase A — Trust and honesty (high impact)", level=2)
    add_bullets(
        doc,
        [
            "Expose ledger_type: canton | simulated | failed on policy APIs and admin UI.",
            "Never show Canton verified when mode=simulated.",
            "Narrow claims verify bypass: require verified=true for canton ledger type.",
            "Optional GCUL_CANTON_STRICT=true — no simulated mint in strict environments.",
            "Health gate: check Canton status before mint; retry queue if offline.",
        ],
    )

    doc.add_heading("Phase B — Correct Daml integration", level=2)
    add_bullets(
        doc,
        [
            "Mint via InsurerMintAuthority.MintPolicy exercise, not raw /v1/create.",
            "Fix payload types (Time, Decimal) for JSON API.",
            "Auto-resolve package ID at startup or from build artifact.",
            "Single idempotent mint path (Pub/Sub or HTTP, not both racing).",
        ],
    )

    doc.add_heading("Phase C — Security and operations", level=2)
    add_bullets(
        doc,
        [
            "Internal API authentication between policy/claims and orchestrator.",
            "Restrict Canton JSON API in cloud (not public unauthenticated).",
            "Signed JWTs or participant credentials outside sandbox.",
            "Monitor Canton health separately from orchestrator.",
        ],
    )

    doc.add_heading("Phase D — Insurance capital market (roadmap)", level=2)
    add_bullets(
        doc,
        [
            "Daml templates: InsuranceLinkedNote, InvestorEligibility, DvP settlement.",
            "Privacy: portfolio-level on-ledger data; no policy-level PII on shared templates.",
            "Reconciliation jobs: Canton contract IDs ↔ policy DB ↔ wallet balances.",
            "Oracle service for catastrophe triggers feeding TriggerLoss choice.",
        ],
    )

    doc.add_heading("4. Quick wins", level=1)
    wins = [
        ["Expose ledger_mode on policy API", "UI and evaluation trail stay truthful"],
        ["Fix assertVerifiedOnCanton for production", "Stops fake Canton attestation"],
        ["Use InsurerMintAuthority in mint path", "Matches Daml design"],
        ["Add Canton to observability", "Ops sees sandbox down early"],
        ["Tokenization banner for simulated mode", "Aligns with enterprise docs"],
    ]
    add_table(doc, ["Change", "Effect"], wins)

    doc.add_heading("5. Key code references", level=1)
    refs = [
        ["canton/daml/daml/Gcul/InsurancePolicy.daml", "InsurerMintAuthority, InsurancePolicy, ClaimSettlement"],
        ["apps/services/blockchain-orchestrator-service/.../CantonJsonApiClient.java", "Canton HTTP mint/verify"],
        ["apps/services/blockchain-orchestrator-service/.../LedgerAdapterRegistry.java", "Canton vs simulated fallback"],
        ["apps/services/claims-service/.../BlockchainValidationClient.java", "Verify bypass logic"],
        ["apps/services/policy-service/.../BlockchainMintClient.java", "Mint API client"],
        ["scripts/local/start-canton.ps1", "Local Docker sandbox"],
        ["deploy/deploy-cloud-run.ps1", "Cloud Canton + orchestrator wiring"],
    ]
    add_table(doc, ["Path", "Role"], refs)

    doc.add_heading("6. Bottom line", level=1)
    doc.add_paragraph(
        "Three structural risks: (1) truth gap — simulated mint looks like Canton mint; "
        "(2) contract gap — Java bypasses authority pattern and type conventions; "
        "(3) production gap — auth, privacy, monitoring, and reconciliation are not yet suitable for LBG-style commercial banking.\n\n"
        "Canton’s value is programmable ownership, eligibility, and atomic DvP — not merely storing policy metadata on a sandbox. "
        "Legal wrapper, SPV, regulated market, and core banking systems remain the source of truth; Canton is the synchronized operational layer."
    )

    doc.add_paragraph()
    foot = doc.add_paragraph("Reboot 2026 Insurance Platform · Internal reference · Not legal or regulatory advice")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    doc = build()
    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
